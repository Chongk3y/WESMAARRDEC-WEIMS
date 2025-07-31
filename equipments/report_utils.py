"""
Utility functions for generating Word document reports from templates
"""
import os
from datetime import datetime
from django.conf import settings
from django.http import HttpResponse
try:
    from docx import Document
    from docx.shared import Inches, Cm, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.enum.section import WD_ORIENT
    from docx.oxml.shared import OxmlElement, qn
except ImportError:
    # Fallback if python-docx is not installed
    Document = None
    WD_ORIENT = None
    Pt = None

# PDF Generation imports
try:
    from reportlab.lib.pagesizes import letter, A4, landscape, portrait
    from reportlab.lib import colors
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import inch
    from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer, Frame, PageTemplate
    from reportlab.platypus.tableofcontents import TableOfContents
    from reportlab.platypus.doctemplate import BaseDocTemplate
    from reportlab.lib.enums import TA_CENTER, TA_LEFT
    from io import BytesIO
except ImportError:
    # Fallback if reportlab is not installed
    SimpleDocTemplate = None
    Table = None


def add_equipment_table_to_doc(doc, equipments, selected_columns, column_labels, table_style='Table Grid', font_size=10):
    """
    Add equipment data table to the Word document with customizable styling
    """
    # Add a paragraph before the table with record count
    p = doc.add_paragraph()
    p.add_run(f"Total Records: {len(equipments)}").bold = True
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add some spacing
    doc.add_paragraph()
    
    # Create table with header row
    table = doc.add_table(rows=1, cols=len(selected_columns))
    
    # Set table style based on user selection
    style_map = {
        'table_grid': 'Table Grid',
        'table_professional': 'Medium Shading 1',
        'table_colorful': 'Colorful Grid',
        'table_minimal': 'Light Grid'
    }
    
    # Try to set the style, fall back to Table Grid if not available
    try:
        table.style = style_map.get(table_style, 'Table Grid')
    except KeyError:
        # If the style doesn't exist, use the default
        table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # Set table to auto-fit
    table.autofit = True
    
    # Add headers
    hdr_cells = table.rows[0].cells
    for i, col in enumerate(selected_columns):
        hdr_cells[i].text = column_labels.get(col, col)
        # Make header bold and set font size
        for paragraph in hdr_cells[i].paragraphs:
            for run in paragraph.runs:
                run.bold = True
                if Pt:  # Only if import was successful
                    run.font.size = Pt(font_size)
                    run.font.name = 'Arial'
    
    # Add data rows
    for eq in equipments:
        row_cells = table.add_row().cells
        for i, col in enumerate(selected_columns):
            value = ""
            if col == 'category':
                value = str(eq.category) if eq.category else ''
            elif col == 'status':
                value = str(eq.status) if eq.status else ''
            elif col == 'created_at':
                value = eq.created_at.strftime('%Y-%m-%d') if eq.created_at else ''
            elif col == 'updated_at':
                value = eq.updated_at.strftime('%Y-%m-%d') if eq.updated_at else ''
            elif col == 'item_purdate':
                value = eq.item_purdate.strftime('%Y-%m-%d') if eq.item_purdate else ''
            elif col == 'date_archived':
                value = eq.date_archived.strftime('%Y-%m-%d') if eq.date_archived else ''
            elif col == 'item_amount':
                value = f"{eq.item_amount:,.2f}" if eq.item_amount else ''
            elif col == 'user_image':
                value = "Image attached" if eq.user_image else "No image"
            else:
                value = str(getattr(eq, col, '')) if getattr(eq, col, '') is not None else ''
            
            row_cells[i].text = value
            # Set font size and style for data cells
            for paragraph in row_cells[i].paragraphs:
                for run in paragraph.runs:
                    if Pt:  # Only if import was successful
                        run.font.size = Pt(font_size)
                        run.font.name = 'Arial'
    
    return table


def replace_placeholders_in_doc(doc, context_data):
    """
    Replace placeholder text in the document with actual data
    Common placeholders: {{DATE}}, {{TOTAL_RECORDS}}, {{REPORT_TITLE}}
    """
    # Replace in paragraphs
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            for placeholder, value in context_data.items():
                if placeholder in run.text:
                    run.text = run.text.replace(placeholder, str(value))
    
    # Replace in tables (headers/footers)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        for placeholder, value in context_data.items():
                            if placeholder in run.text:
                                run.text = run.text.replace(placeholder, str(value))
    
    # Replace in headers and footers
    for section in doc.sections:
        # Header
        if section.header:
            for paragraph in section.header.paragraphs:
                for run in paragraph.runs:
                    for placeholder, value in context_data.items():
                        if placeholder in run.text:
                            run.text = run.text.replace(placeholder, str(value))
        
        # Footer
        if section.footer:
            for paragraph in section.footer.paragraphs:
                for run in paragraph.runs:
                    for placeholder, value in context_data.items():
                        if placeholder in run.text:
                            run.text = run.text.replace(placeholder, str(value))


def generate_word_report_from_template(equipments, selected_columns, column_labels, filename="equipment_report.docx", 
                                      orientation="portrait", table_style="table_grid", font_size=10):
    """
    Generate a Word document report using WESMAARRDEC template with customizable options
    """
    if Document is None:
        raise ImportError("python-docx is required for Word document generation")
    
    # Load the WESMAARRDEC template
    template_path = os.path.join(settings.MEDIA_ROOT, 'report_templates', 'WESMAARRDEC template.docx')
    
    if not os.path.exists(template_path):
        raise FileNotFoundError(f"WESMAARRDEC template not found at: {template_path}")
    
    # Load template document
    doc = Document(template_path)
    
    # Set page orientation
    if WD_ORIENT:  # Only if import was successful
        for section in doc.sections:
            if orientation == "landscape":
                # Switch to landscape
                new_width, new_height = section.page_height, section.page_width
                section.orientation = WD_ORIENT.LANDSCAPE
                section.page_width = new_width
                section.page_height = new_height
            else:
                # Ensure portrait
                section.orientation = WD_ORIENT.PORTRAIT
    
    # Prepare context data for placeholder replacement
    context_data = {
        '{{DATE}}': datetime.now().strftime('%B %d, %Y - %I:%M %p'),
        '{{CURRENT_DATE}}': datetime.now().strftime('%B %d, %Y'),
        '{{CURRENT_TIME}}': datetime.now().strftime('%I:%M %p'),
        '{{TOTAL_RECORDS}}': str(len(equipments)),
        '{{REPORT_TITLE}}': 'Equipment Inventory Report',
        '{{GENERATED_BY}}': 'WESMAARRDEC System',
        '{{RECORD_COUNT}}': str(len(equipments)),
        '{{ORIENTATION}}': orientation.title(),
    }
    
    # Replace placeholders in headers and footers first
    replace_placeholders_in_doc(doc, context_data)
    
    # Clear existing content in the document body (keep headers/footers)
    # Remove all existing paragraphs except headers/footers
    for paragraph in doc.paragraphs[:]:
        p = paragraph._element
        p.getparent().remove(p)
    
    # Add report title and date to document body
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_para.add_run("Equipment Inventory Report")
    title_run.font.size = Pt(18)
    title_run.font.bold = True
    title_run.font.name = 'Arial'
    
    # Add date and info section
    info_para = doc.add_paragraph()
    info_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_run = info_para.add_run(f"Generated on: {datetime.now().strftime('%B %d, %Y at %I:%M %p')}")
    date_run.font.size = Pt(12)
    date_run.font.name = 'Arial'
    
    # Add spacing before table
    doc.add_paragraph()
    
    # Add the equipment table with custom styling
    add_equipment_table_to_doc(doc, equipments, selected_columns, column_labels, table_style, font_size)
    
    return doc


def create_word_response(doc, filename="equipment_report.docx"):
    """
    Create HttpResponse with Word document for download
    """
    # Ensure filename has .docx extension
    if not filename.endswith('.docx'):
        filename = filename.replace('.xlsx', '.docx').replace('.xls', '.docx')
        if not filename.endswith('.docx'):
            filename += '.docx'
    
    response = HttpResponse(
        content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )
    # Use proper quotes around filename to prevent Excel interpretation
    response['Content-Disposition'] = f'attachment; filename="{filename}"'
    
    # Force download behavior and prevent Excel interpretation
    response['Cache-Control'] = 'no-cache, no-store, must-revalidate'
    response['Pragma'] = 'no-cache'
    response['Expires'] = '0'
    response['Content-Description'] = 'File Transfer'
    response['Content-Transfer-Encoding'] = 'binary'
    
    # Debug output
    print(f"DEBUG: Creating Word response with content-type: application/vnd.openxmlformats-officedocument.wordprocessingml.document")
    print(f"DEBUG: Filename: {filename}")
    
    # Save document to response
    doc.save(response)
    
    print(f"DEBUG: Word document saved to response successfully")
    return response


# ==================== PDF GENERATION FUNCTIONS ====================

def extract_template_styling(template_path):
    """
    Extract header and footer content from Word template for use in PDF generation
    """
    try:
        doc = Document(template_path)
        template_info = {
            'header_text': '',
            'footer_text': '',
            'has_header': False,
            'has_footer': False
        }
        
        # Extract header information
        for section in doc.sections:
            if section.header:
                template_info['has_header'] = True
                for para in section.header.paragraphs:
                    if para.text.strip():
                        template_info['header_text'] += para.text.strip() + '\n'
                    # Check for runs even if paragraph text is empty (in case of formatting)
                    elif any(run.text.strip() for run in para.runs):
                        for run in para.runs:
                            if run.text.strip():
                                template_info['header_text'] += run.text.strip() + ' '
                        template_info['header_text'] += '\n'
            
            if section.footer:
                template_info['has_footer'] = True
                for para in section.footer.paragraphs:
                    if para.text.strip():
                        template_info['footer_text'] += para.text.strip() + '\n'
                    # Check for runs even if paragraph text is empty
                    elif any(run.text.strip() for run in para.runs):
                        for run in para.runs:
                            if run.text.strip():
                                template_info['footer_text'] += run.text.strip() + ' '
                        template_info['footer_text'] += '\n'
        
        # Clean up trailing newlines and spaces
        template_info['header_text'] = template_info['header_text'].strip()
        template_info['footer_text'] = template_info['footer_text'].strip()
        
        # If no text found but header/footer exist, use defaults
        if template_info['has_header'] and not template_info['header_text']:
            template_info['header_text'] = 'WESMAARRDEC\nEquipment Inventory Management System'
        
        if template_info['has_footer'] and not template_info['footer_text']:
            template_info['footer_text'] = 'WESMAARRDEC Equipment Report'
        
        return template_info
    
    except Exception as e:
        # If template can't be read, return default info
        return {
            'header_text': 'WESMAARRDEC\nEquipment Inventory Management System',
            'footer_text': 'WESMAARRDEC Equipment Report',
            'has_header': True,
            'has_footer': True
        }


def generate_pdf_report_from_template(equipments, selected_columns, column_labels, 
                                     filename="equipment_report.pdf", orientation="portrait", 
                                     table_style="table_grid", font_size=10):
    """
    Generate PDF report using styling extracted from WESMAARRDEC Word template with header/footer
    """
    if SimpleDocTemplate is None:
        raise ImportError("ReportLab is not installed. Cannot generate PDF documents.")
    
    # Extract styling from Word template
    template_path = os.path.join(settings.MEDIA_ROOT, 'report_templates', 'WESMAARRDEC template.docx')
    template_info = extract_template_styling(template_path)
    
    # Create a BytesIO buffer to hold the PDF
    buffer = BytesIO()
    
    # Set page size based on orientation
    if orientation == "landscape":
        pagesize = landscape(A4)
    else:
        pagesize = portrait(A4)

    # Custom document template class with header/footer
    class WESMDocTemplate(BaseDocTemplate):
        """Custom document template with header and footer from Word template"""
        
        def __init__(self, filename, **kwargs):
            self.template_info = template_info
            BaseDocTemplate.__init__(self, filename, **kwargs)
            
        def draw_header_footer(self, canvas, doc):
            """Draw header and footer on each page using template content"""
            canvas.saveState()
            
            # Get page dimensions
            page_width, page_height = self.pagesize
            
            # Draw header if available from template
            if self.template_info['has_header'] and self.template_info['header_text']:
                # Clean header text (remove extra whitespace and newlines)
                header_text = ' '.join(self.template_info['header_text'].split())
                
                # Header styling - make it prominent like the Word template
                canvas.setFont("Helvetica-Bold", font_size + 4)
                canvas.setFillColor(colors.black)
                
                # Center the header text
                text_width = canvas.stringWidth(header_text, "Helvetica-Bold", font_size + 4)
                x_position = (page_width - text_width) / 2
                y_position = page_height - 0.6*inch
                
                canvas.drawString(x_position, y_position, header_text)
                
                # Add decorative line under header
                canvas.setStrokeColor(colors.darkblue)
                canvas.setLineWidth(2)
                canvas.line(0.75*inch, page_height - 0.8*inch, 
                           page_width - 0.75*inch, page_height - 0.8*inch)
            
            # Draw footer if available from template
            if self.template_info['has_footer'] and self.template_info['footer_text']:
                # Clean footer text (remove extra whitespace and newlines)
                footer_text = ' '.join(self.template_info['footer_text'].split())
                
                # Footer styling - smaller and gray like typical document footers
                canvas.setFont("Helvetica", font_size - 1)
                canvas.setFillColor(colors.grey)
                
                # Center the footer text
                text_width = canvas.stringWidth(footer_text, "Helvetica", font_size - 1)
                x_position = (page_width - text_width) / 2
                y_position = 0.6*inch
                
                canvas.drawString(x_position, y_position, footer_text)
                
                # Add subtle line above footer
                canvas.setStrokeColor(colors.lightgrey)
                canvas.setLineWidth(1)
                canvas.line(0.75*inch, 0.8*inch, 
                           page_width - 0.75*inch, 0.8*inch)
            
            # Add page numbers and generation info in footer area
            from datetime import datetime
            page_num = canvas.getPageNumber()
            canvas.setFont("Helvetica", font_size - 2)
            canvas.setFillColor(colors.grey)
            
            # Page number on bottom right
            canvas.drawRightString(page_width - 0.75*inch, 0.3*inch, f"Page {page_num}")
            
            # Generation date on bottom left
            gen_date = datetime.now().strftime("%Y-%m-%d %H:%M")
            canvas.drawString(0.75*inch, 0.3*inch, f"Generated: {gen_date}")
            
            canvas.restoreState()

    # Create custom document with enhanced header/footer support
    doc = WESMDocTemplate(
        buffer,
        pagesize=pagesize,
        rightMargin=0.75*inch,
        leftMargin=0.75*inch,
        topMargin=1.4*inch,  # Increased for header with decorative line
        bottomMargin=1.2*inch  # Increased for footer with page info
    )
    
    # Create frame for content (accounting for header/footer space)
    from reportlab.platypus import Frame, PageTemplate
    frame = Frame(
        0.75*inch, 1.2*inch,
        pagesize[0] - 1.5*inch, pagesize[1] - 2.6*inch,
        id='normal',
        showBoundary=0
    )
    
    # Create page template with our custom header/footer function
    template = PageTemplate(id='main', frames=[frame], onPage=doc.draw_header_footer)
    doc.addPageTemplates([template])

    # Container for the 'Flowable' objects
    story = []
    
    # Define styles
    styles = getSampleStyleSheet()
    
    # Custom report title style (main content title, not header)
    report_title_style = ParagraphStyle(
        'WESMReportTitle',
        parent=styles['Heading1'],
        fontSize=font_size + 6,
        spaceAfter=20,
        alignment=TA_CENTER,
        textColor=colors.darkblue,
        fontName='Helvetica-Bold'
    )
    
    # Custom subtitle style
    subtitle_style = ParagraphStyle(
        'WESMSubtitle',
        parent=styles['Normal'],
        fontSize=font_size + 2,
        spaceAfter=15,
        alignment=TA_CENTER,
        textColor=colors.black,
        fontName='Helvetica'
    )
    
    # Custom date style
    date_style = ParagraphStyle(
        'WESMDate',
        parent=styles['Normal'],
        fontSize=font_size,
        spaceAfter=12,
        alignment=TA_CENTER,
        textColor=colors.grey,
        fontName='Helvetica'  # Fixed: Use standard Helvetica instead of Helvetica-Italic
    )
    
    # Add main report title (header content is now handled by page template)
    report_title = Paragraph("Equipment Inventory Report", report_title_style)
    story.append(report_title)
    
    # Add date and spacing
    current_date = datetime.now().strftime("%B %d, %Y at %I:%M %p")
    date_para = Paragraph(f"Generated on: {current_date}", date_style)
    story.append(date_para)
    
    # Add spacing before table
    story.append(Spacer(1, 20))
    
    # Add record count
    record_count = Paragraph(f"Equipment Report - {len(equipments)} records", styles['Heading3'])
    story.append(record_count)
    story.append(Spacer(1, 10))
    
    # Prepare table data
    table_data = []
    
    # Add header row
    headers = [column_labels.get(col, col) for col in selected_columns]
    table_data.append(headers)
    
    # Add data rows
    for eq in equipments:
        row = []
        for col in selected_columns:
            if col == 'category':
                value = str(eq.category) if eq.category else ''
            elif col == 'status':
                value = str(eq.status) if eq.status else ''
            elif col == 'created_at':
                value = eq.created_at.strftime('%Y-%m-%d') if eq.created_at else ''
            elif col == 'updated_at':
                value = eq.updated_at.strftime('%Y-%m-%d') if eq.updated_at else ''
            elif col == 'item_amount':
                value = f"{eq.item_amount:.2f}" if eq.item_amount else ''
            elif col == 'item_purdate':
                value = eq.item_purdate.strftime('%Y-%m-%d') if eq.item_purdate else ''
            elif col == 'date_archived':
                value = eq.date_archived.strftime('%Y-%m-%d') if eq.date_archived else ''
            elif col == 'is_returned':
                value = 'Yes' if eq.is_returned else 'No'
            elif col == 'is_archived':
                value = 'Yes' if eq.is_archived else 'No'
            else:
                value = str(getattr(eq, col, '')) if getattr(eq, col, '') is not None else ''
            
            # Truncate long text to prevent table overflow
            if len(str(value)) > 30:
                value = str(value)[:27] + "..."
            
            row.append(str(value))
        table_data.append(row)
    
    # Create table
    if table_data:
        # Calculate column widths based on page size and number of columns
        page_width = pagesize[0] - 1*inch  # Account for margins
        col_width = page_width / len(selected_columns)
        
        # Ensure minimum column width
        if col_width < 0.8*inch:
            col_width = 0.8*inch
        
        col_widths = [col_width] * len(selected_columns)
        
        table = Table(table_data, colWidths=col_widths)
        
        # Apply table style based on user selection
        if table_style == "table_professional":
            table_style_commands = [
                ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('FONTSIZE', (0, 0), (-1, 0), font_size + 1),
                ('FONTSIZE', (0, 1), (-1, -1), font_size),
                ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
                ('GRID', (0, 0), (-1, -1), 1, colors.black)
            ]
        elif table_style == "table_colorful":
            table_style_commands = [
                ('BACKGROUND', (0, 0), (-1, 0), colors.blue),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
                ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('FONTSIZE', (0, 0), (-1, 0), font_size + 1),
                ('FONTSIZE', (0, 1), (-1, -1), font_size),
                ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                ('BACKGROUND', (0, 1), (-1, -1), colors.lightblue),
                ('GRID', (0, 0), (-1, -1), 1, colors.darkblue)
            ]
        elif table_style == "table_minimal":
            table_style_commands = [
                ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('FONTSIZE', (0, 0), (-1, 0), font_size + 1),
                ('FONTSIZE', (0, 1), (-1, -1), font_size),
                ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                ('LINEBELOW', (0, 0), (-1, 0), 2, colors.black),
                ('LINEBELOW', (0, -1), (-1, -1), 1, colors.grey)
            ]
        else:  # table_grid (default)
            table_style_commands = [
                ('BACKGROUND', (0, 0), (-1, 0), colors.lightgrey),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.black),
                ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('FONTSIZE', (0, 0), (-1, 0), font_size + 1),
                ('FONTSIZE', (0, 1), (-1, -1), font_size),
                ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                ('GRID', (0, 0), (-1, -1), 1, colors.black)
            ]
        
        table.setStyle(TableStyle(table_style_commands))
        story.append(table)
    
    # Footer content is now handled by the page template's draw_header_footer method
    
    # Build PDF with template-based header/footer
    doc.build(story)
    
    # Get the value of the BytesIO buffer and return it
    pdf_content = buffer.getvalue()
    buffer.close()
    
    return pdf_content


def create_pdf_response(pdf_content, filename="equipment_report.pdf"):
    """
    Create HttpResponse with PDF document for download
    """
    response = HttpResponse(content_type='application/pdf')
    response['Content-Disposition'] = f'attachment; filename="{filename}"'
    # Force download behavior
    response['Cache-Control'] = 'no-cache, no-store, must-revalidate'
    response['Pragma'] = 'no-cache'
    response['Expires'] = '0'
    
    response.write(pdf_content)
    return response
